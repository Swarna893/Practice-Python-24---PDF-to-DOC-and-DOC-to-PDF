# Practice-Python-24---PDF-to-DOC-and-DOC-to-PDF

# CONVERT PDF TO WORD
'''
import fitz
from docx import Document


def pdf_to_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text()
    return text


def text_to_word(text, output_path):
    doc = Document()
    doc.add_paragraph(text)

    # Save the Word document
    doc.save(output_path)
    print(f"Conversion successful. Word document saved at {output_path}")


if __name__ == "__main__":
    pdf_path = "Policy_Certificate_49982828_25122023_policyPDF.pdf"  # Replace with the path to your PDF file
    output_path = "output.docx"  # Replace with the desired output path

    pdf_text = pdf_to_text(pdf_path)
    text_to_word(pdf_text, output_path)
'''


# CONVERT WORD TO PDF

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter


def convert_docx_to_pdf(docx_filename, pdf_filename):
    # Read the Word document
    doc = Document(docx_filename)

    # Create a PDF file
    pdf_file = canvas.Canvas(pdf_filename, pagesize=letter)

    for para in doc.paragraphs:
        # Write each paragraph to the PDF file
        pdf_file.drawString(100, 800, para.text)
        # You may need to adjust the coordinates based on your layout

    # Save the PDF file
    pdf_file.save()


if __name__ == "__main__":
    # Specify the input Word document and output PDF file
    input_docx = "Jyoti Updated CV JOB.docx"
    output_pdf = "output_document.pdf"

    # Convert Word to PDF
    convert_docx_to_pdf(input_docx, output_pdf)

    print(f"Conversion complete. PDF file saved as {output_pdf}")



